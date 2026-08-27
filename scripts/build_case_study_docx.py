from __future__ import annotations

from pathlib import Path
from shutil import copy2

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "assets"
OUTPUT = ROOT / "deliverables" / "CareMatch_Modern_Data_Stack_Case_Study.docx"
SOURCE_ARCH = Path(
    r"C:\Users\mindo\AppData\Local\Temp\codex-clipboard-96c81feb-8ba2-4b26-8b70-fe5b8c30ee2d.png"
)

NAVY = "17365D"
BLUE = "2E75B6"
PALE_BLUE = "DDEBF7"
GREEN = "2E7D32"
PALE_GREEN = "E2F0D9"
AMBER = "9C6500"
PALE_AMBER = "FFF2CC"
GRAY = "666666"
PALE_GRAY = "F2F4F7"
RED = "A61B1B"
WHITE = "FFFFFF"
BLACK = "111111"
LIST_COUNTER = 100


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent=120) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, size=10.5, bold=False, color=BLACK, italic=False, font="Aptos") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def add_paragraph(doc, text: str, *, bold_lead: str | None = None, after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run(rest, italic=italic)
    else:
        run = p.add_run(text)
        set_run(run, italic=italic)
    return p


def add_list(doc, items: list[str], ordered: bool) -> None:
    global LIST_COUNTER
    LIST_COUNTER += 1
    numbering = doc.part.numbering_part.element
    abstract_id = LIST_COUNTER
    num_id = LIST_COUNTER

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for text in items:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p_pr.append(num_pr)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        set_run(p.add_run(text))


def add_bullets(doc, items: list[str]) -> None:
    add_list(doc, items, ordered=False)


def add_numbered(doc, items: list[str]) -> None:
    add_list(doc, items, ordered=True)


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_code(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_cell_like_paragraph(p, "F7F7F7")
    run = p.add_run(text)
    set_run(run, size=8.5, color="222222", font="Consolas")


def set_cell_like_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "D9D9D9")
        p_bdr.append(el)
    p_pr.append(p_bdr)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    tr_pr.append(table_header)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(header), size=9, bold=True, color=WHITE)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(str(value)), size=8.8)
            if status_col is not None and index == status_col:
                lower = str(value).lower()
                if "verified" in lower or "complete" in lower or "passed" in lower:
                    set_cell_shading(cells[index], PALE_GREEN)
                elif "blocked" in lower or "required" in lower or "pending" in lower:
                    set_cell_shading(cells[index], PALE_AMBER)
                else:
                    set_cell_shading(cells[index], PALE_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label: str, text: str, fill=PALE_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(f"{label}: "), bold=True, color=color)
    set_run(p.add_run(text), color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc, path: Path, caption: str, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(9)
    set_run(cap.add_run(caption), size=8.5, italic=True, color=GRAY)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(paragraph.add_run("CareMatch case study  |  "), size=8.5, color=GRAY)
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def create_visuals() -> tuple[Path, Path, Path, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    arch_reference = ASSETS / "intelycare_reference_architecture.png"
    if SOURCE_ARCH.exists():
        copy2(SOURCE_ARCH, arch_reference)

    regular_path = Path(r"C:\Windows\Fonts\arial.ttf")
    bold_path = Path(r"C:\Windows\Fonts\arialbd.ttf")

    def font(size: int, bold: bool = False):
        path = bold_path if bold else regular_path
        return ImageFont.truetype(str(path), size) if path.exists() else ImageFont.load_default()

    def centered(draw, area, text, fnt, fill="#111111"):
        left, top, right, bottom = area
        lines = text.split("\n")
        heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
        y = top + ((bottom - top) - sum(heights) - 6 * (len(lines) - 1)) / 2
        for line, height in zip(lines, heights):
            box = draw.textbbox((0, 0), line, font=fnt)
            x = left + ((right - left) - (box[2] - box[0])) / 2
            draw.text((x, y), line, font=fnt, fill=fill)
            y += height + 6

    def arrow(draw, start, end, color="#4472C4", width=5):
        draw.line([start, end], fill=color, width=width)
        x, y = end
        if abs(end[0] - start[0]) >= abs(end[1] - start[1]):
            points = [(x, y), (x - 18, y - 10), (x - 18, y + 10)]
        else:
            points = [(x, y), (x - 10, y - 18), (x + 10, y - 18)]
        draw.polygon(points, fill=color)

    actual = ASSETS / "carematch_actual_and_target_architecture.png"
    image = Image.new("RGB", (2400, 1250), "white")
    draw = ImageDraw.Draw(image)
    draw.text((80, 55), "CareMatch architecture: verified core and account-owned connector work", font=font(48, True), fill="#17365D")
    palette = {"verified": ("#E2F0D9", "#2E7D32"), "pending": ("#FFF2CC", "#9C6500"), "planned": ("#F2F4F7", "#666666")}

    def box(area, label, state="verified"):
        fill, edge = palette[state]
        draw.rounded_rectangle(area, radius=24, fill=fill, outline=edge, width=5)
        centered(draw, area, label, font(24, True))

    boxes = [
        ((70, 300, 370, 480), "Six synthetic\nsource families", "verified"),
        ((430, 300, 680, 480), "EC2\nAirflow", "verified"),
        ((740, 300, 990, 480), "Amazon S3\nraw lake", "verified"),
        ((1050, 300, 1300, 480), "Snowflake\nRAW", "verified"),
        ((1360, 300, 1660, 480), "dbt\nSTAGING + ANALYTICS", "verified"),
        ((1720, 300, 1970, 480), "Hightouch\nsource", "verified"),
        ((1720, 600, 1970, 780), "Slack OAuth\nand live sync", "pending"),
        ((1360, 900, 1660, 1080), "Fivetran\ndestination", "verified"),
        ((1000, 900, 1300, 1080), "Marketo, Pendo,\nSurveyMonkey", "pending"),
        ((1720, 900, 2290, 1080), "Salesforce, OneDrive,\nGoogle Ads, Meta", "planned"),
    ]
    for item in boxes:
        box(*item)
    for start, end in [((370, 390), (430, 390)), ((680, 390), (740, 390)), ((990, 390), (1050, 390)), ((1300, 390), (1360, 390)), ((1660, 390), (1720, 390)), ((1845, 480), (1845, 600)), ((1300, 990), (1360, 990)), ((1510, 900), (1200, 480)), ((1845, 780), (1845, 900))]:
        arrow(draw, start, end)
    draw.text((80, 1160), "Green: live or tested     Amber: external authorization or subscription required     Gray: planned scope", font=font(24), fill="#444444")
    image.save(actual)

    growth = ASSETS / "incremental_nurse_growth.png"
    image = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((80, 45), "Five daily snapshots grow RAW history while staging stays current", font=font(42, True), fill="#17365D")
    left, top, right, bottom = 170, 180, 1690, 740
    for value in range(0, 2501, 500):
        y = bottom - int(value / 2500 * (bottom - top))
        draw.line((left, y, right, y), fill="#D9E2F3", width=2)
        draw.text((65, y - 13), f"{value:,}", font=font(22), fill="#555555")
    dates = ["23 Aug", "24 Aug", "25 Aug", "26 Aug", "27 Aug"]
    raw = [500, 1000, 1500, 2000, 2500]
    staging = [500] * 5
    xs = [left + int(i * (right - left) / 4) for i in range(5)]
    raw_points = [(x, bottom - int(y / 2500 * (bottom - top))) for x, y in zip(xs, raw)]
    stage_points = [(x, bottom - int(y / 2500 * (bottom - top))) for x, y in zip(xs, staging)]
    draw.line(raw_points, fill="#2E75B6", width=8)
    draw.line(stage_points, fill="#2E7D32", width=8)
    for x, y in raw_points:
        draw.ellipse((x - 10, y - 10, x + 10, y + 10), fill="#2E75B6")
    for x, y in stage_points:
        draw.ellipse((x - 9, y - 9, x + 9, y + 9), fill="#2E7D32")
    for x, label, value, point in zip(xs, dates, raw, raw_points):
        draw.text((x - 45, bottom + 25), label, font=font(22), fill="#333333")
        draw.text((x - 32, point[1] - 45), f"{value:,}", font=font(20, True), fill="#2E75B6")
    draw.line((230, 820, 300, 820), fill="#2E75B6", width=8)
    draw.text((320, 803), "RAW nurse rows", font=font(23), fill="#333333")
    draw.line((760, 820, 830, 820), fill="#2E7D32", width=8)
    draw.text((850, 803), "Current staging nurses", font=font(23), fill="#333333")
    image.save(growth)

    evidence = ASSETS / "verification_dashboard.png"
    image = Image.new("RGB", (2000, 980), "white")
    draw = ImageDraw.Draw(image)
    draw.text((75, 55), "Final audit evidence - 27 August 2026", font=font(48, True), fill="#17365D")
    cards = [
        (70, 210, "5", "dated S3 batches"), (555, 210, "55", "raw S3 objects"),
        (1040, 210, "8 / 8", "local tests passed"), (1525, 210, "200", "Airflow health response"),
        (70, 590, "11", "Snowflake RAW tables"), (555, 590, "2,500", "RAW nurse snapshots"),
        (1040, 590, "500", "current nurses"), (1525, 590, "0", "duplicate nurses in mart"),
    ]
    for x, y, metric, label in cards:
        area = (x, y, x + 400, y + 280)
        draw.rounded_rectangle(area, radius=24, fill="#F2F4F7", outline="#B4C7E7", width=5)
        centered(draw, (x, y + 25, x + 400, y + 165), metric, font(55, True), fill="#2E75B6")
        centered(draw, (x + 10, y + 155, x + 390, y + 255), label, font(23), fill="#333333")
    image.save(evidence)

    return arch_reference, actual, growth, evidence


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for level, size, before, after, color in [
        (1, 16, 16, 8, BLUE),
        (2, 13, 12, 6, BLUE),
        (3, 11.5, 8, 4, NAVY),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("CAREMATCH MODERN DATA STACK"), size=8.5, bold=True, color=GRAY)
    add_page_number(section.footer.paragraphs[0])


def build() -> None:
    arch_reference, actual, growth, evidence = create_visuals()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    cover = doc.add_paragraph()
    cover.paragraph_format.space_before = Pt(72)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(cover.add_run("CAREMATCH"), size=12, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(10)
    set_run(title.add_run("Modern Data Stack Case Study"), size=28, bold=True, color=NAVY, font="Aptos Display")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_run(subtitle.add_run("A small scale implementation inspired by the IntelyCare architecture"), size=13, color=GRAY)
    add_figure(doc, evidence, "Verified implementation evidence as of 27 August 2026", width=6.2)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(16)
    set_run(meta.add_run("Prepared for technical review, business demonstration and system design discussion"), size=10, color=GRAY)
    doc.add_page_break()

    add_heading(doc, "Executive summary")
    add_paragraph(doc, "CareMatch is a small scale healthcare staffing data platform built to reproduce the main ideas in the IntelyCare customer story. It creates synthetic workforce and engagement data on an Airflow server in Amazon EC2, lands dated files in Amazon S3, loads them into Snowflake and transforms them with dbt. Hightouch and Fivetran service connections are prepared for activation and managed ingestion.")
    add_callout(doc, "Current position", "The core path from EC2 Airflow to S3, Snowflake and dbt is live and tested. Hightouch can read Snowflake and Fivetran can write to Snowflake. Slack, Marketo, Pendo and SurveyMonkey still need account owned authorization before any external data movement can be claimed.", fill=PALE_GREEN, color=GREEN)
    add_paragraph(doc, "The platform now contains five dated source batches. S3 holds 55 raw objects. Snowflake holds 2,500 nurse snapshot rows in RAW while dbt keeps 500 current nurses in staging. This difference is intentional. It shows append only history followed by business key deduplication.")
    add_paragraph(doc, "The design is suitable for a case study, a classroom demonstration and a small internal proof. It is not yet a highly available production platform. The final section explains what must change before production use.")

    add_heading(doc, "What was taken from the IntelyCare story")
    add_paragraph(doc, "The supplied Snowflake case study describes IntelyCare as a technology enabled nurse staffing platform. It explains a move away from fragmented operational databases toward Snowflake, dbt, Fivetran and Hightouch. It also describes the need to match nurses with shifts, react to changing demand and activate trusted data in customer facing tools.")
    add_table(doc, ["Original story theme", "CareMatch implementation"], [
        ["Central analytics platform", "Snowflake with RAW, STAGING and ANALYTICS schemas"],
        ["Reliable data ingestion", "Airflow to S3 plus a prepared Fivetran landing database"],
        ["Tested transformation layer", "dbt staging models, marts and data tests"],
        ["Business activation", "A governed at risk nurse audience prepared for Hightouch"],
        ["Operational resilience", "Repeatable Terraform, manifests, checksums and idempotent loads"],
    ], [3000, 6360])
    add_callout(doc, "Important boundary", "The financial and growth results in the original PDF belong to IntelyCare. They are background evidence for the architecture. They are not results produced by this synthetic CareMatch implementation.", fill=PALE_AMBER, color=AMBER)
    if arch_reference.exists():
        add_figure(doc, arch_reference, "Reference architecture supplied for the case study", width=6.45)

    add_heading(doc, "Business problem and intended users")
    add_paragraph(doc, "Healthcare staffing teams need one trusted view of nurse supply, open shift demand, applications, assignments, safety checks, marketing performance and product engagement. Without a shared data layer each department receives a different answer and operational teams rely on manual files.")
    add_bullets(doc, [
        "Operations needs to identify open shifts and qualified nurses in each market.",
        "Marketing needs to measure spend against applicants and completed assignments.",
        "Product teams need app events and model scores for engagement decisions.",
        "Sales and support need approved audience fields without direct access to raw data.",
        "Data teams need repeatable loads, clear ownership and evidence when a run fails.",
    ])

    add_heading(doc, "System architecture")
    add_figure(doc, actual, "Verified core, external authorization boundary and planned connector scope", width=6.45)
    add_heading(doc, "Component responsibilities", 2)
    add_table(doc, ["Component", "Responsibility", "Why it is here"], [
        ["Python generator", "Creates deterministic synthetic healthcare data", "Safe repeatable source data without real patient or employee records"],
        ["Airflow on EC2", "Schedules generation and six parallel upload branches", "Visible orchestration at low proof of concept cost"],
        ["Amazon S3", "Stores immutable dated source files and manifests", "Durable replay point between compute and warehouse"],
        ["Snowflake RAW", "Loads 11 source entities from S3", "Append only audit history with file load protection"],
        ["dbt", "Cleans, types, deduplicates and builds business models", "Version controlled SQL and tests"],
        ["Hightouch", "Reads approved analytics models and sends changes", "Reverse ETL without custom destination code"],
        ["Fivetran", "Pulls SaaS data into a separate landing database", "Managed API maintenance and connector cursors"],
    ], [1600, 3960, 3800])

    add_heading(doc, "Data sources and current volume")
    add_paragraph(doc, "The generator produces six source families. Fourteen files are created locally because some domains contain more than one file. Snowflake creates 11 RAW tables because two event files share the APP EVENTS table and several operational files map into the same source family rather than a separate family table count.")
    add_table(doc, ["Source family", "Main entities", "Format", "Typical rows per batch"], [
        ["Operational", "Facilities, nurses, shifts, applications, assignments, health screenings", "CSV", "about 16,000"],
        ["External", "Market conditions", "CSV", "15"],
        ["Data science", "Nurse scores", "CSV", "500"],
        ["Appcast", "Campaign performance", "CSV", "4"],
        ["App stream", "Product events", "JSON Lines", "about 3,000"],
        ["Spreadsheets", "Manual overrides", "CSV", "about 20"],
    ], [1600, 3900, 1300, 2560])
    add_paragraph(doc, "Five batches dated 23 August through 27 August 2026 are present in S3. They contain 55 raw objects with a combined size of about 8.6 MB. Five manifests are stored separately.")

    add_heading(doc, "Initial load design")
    add_paragraph(doc, "The first load creates the baseline. It is intentionally simple so every stage can be explained and verified.")
    add_numbered(doc, [
        "Airflow chooses the scheduled date or a date supplied in the manual run configuration.",
        "The generator creates deterministic rows and a manifest with schemas, row counts, byte counts and SHA 256 checksums.",
        "Six Airflow upload tasks run in parallel and write files to source, entity and load date paths in S3.",
        "A final manifest is written only after every source family succeeds.",
        "Snowflake COPY statements load each entity into its RAW table.",
        "dbt builds five staging views and five analytics tables.",
        "Data tests check keys, relationships and audience consent rules.",
    ])
    add_code(doc, "raw/source=operational/entity=nurses/load_date=2026-08-23/nurses.csv\nmanifests/load_date=2026-08-23/manifest.json")
    add_callout(doc, "Initial load result", "The first dated partition established 500 nurses, 3,000 shifts and the related activity needed for the business models. The source files remain available in S3 for replay.")

    add_heading(doc, "Incremental loading design")
    add_figure(doc, growth, "RAW keeps all five daily snapshots while staging keeps one current nurse record", width=6.25)
    add_paragraph(doc, "Incremental behavior exists at several layers. Airflow creates a new date partition. S3 keeps the new files beside the old files. Snowflake COPY history skips filenames that were already loaded. dbt staging uses a row number window over each business key and keeps the newest record. Hightouch is designed to compare model results by primary key. Fivetran uses connector managed cursors after the initial sync.")
    add_table(doc, ["Layer", "Initial behavior", "Later behavior", "Duplicate control"], [
        ["Airflow", "Creates a complete baseline batch", "Creates a new dated batch", "One active run and deterministic date seed"],
        ["S3", "Stores the first set of objects", "Adds a new load date path", "Partitioned object names and versioning"],
        ["Snowflake", "Loads unseen files", "Loads only new filenames", "COPY load metadata, no FORCE option"],
        ["dbt staging", "Creates the current state", "Selects the newest business record", "Row number by stable key"],
        ["Hightouch", "Creates a CDC baseline", "Sends added, changed or removed rows", "NURSE ID primary key"],
        ["Fivetran", "Performs historical source sync", "Uses source specific cursors", "Connector state and primary keys"],
    ], [1300, 2300, 3000, 2760])
    add_heading(doc, "How to prove incremental loading live", 2)
    add_numbered(doc, [
        "Record the RAW nurse count in Snowflake.",
        "Trigger the Airflow DAG with a date that does not exist in S3.",
        "Show eight successful tasks and the new S3 manifest.",
        "Run the Snowflake load SQL once and show that RAW increases by 500 nurse rows.",
        "Run the same load SQL again and show a zero row change.",
        "Run dbt and show that staging still contains 500 unique nurses.",
    ])
    add_code(doc, "SELECT TO_DATE(RECORD_UPDATED_AT) AS BATCH_DATE,\n       COUNT(*) AS ROW_COUNT,\n       COUNT(DISTINCT NURSE_ID) AS UNIQUE_NURSES\nFROM CAREMATCH.RAW.NURSES\nGROUP BY 1\nORDER BY 1;")
    add_callout(doc, "Observed result", "RAW contains 2,500 nurse rows across five dates. STAGING contains 500 rows and 500 unique nurse IDs. The current state is stable while audit history grows.", fill=PALE_GREEN, color=GREEN)

    add_heading(doc, "Transformation logic")
    add_heading(doc, "Staging layer", 2)
    add_bullets(doc, [
        "Nurses, shifts, applications and assignments are typed and standardized.",
        "Records are ranked inside each business key and the latest version is selected.",
        "Application events are read from Snowflake VARIANT and converted into typed columns.",
        "Stable keys remain available for relationships and downstream change detection.",
    ])
    add_heading(doc, "Analytics layer", 2)
    add_table(doc, ["Model", "Purpose", "Materialization"], [
        ["DIM NURSES", "Current nurse profile and engagement attributes", "Table"],
        ["FCT SHIFT PERFORMANCE", "Shift demand, assignment and outcome measures", "Table"],
        ["MART MARKETING EFFICIENCY", "Spend, applicants and quality by campaign", "Table"],
        ["MART MARKET SUPPLY DEMAND", "Demand, supply and rate comparison by market", "Table"],
        ["AUDIENCE AT RISK NURSES", "Consent safe activation list for Hightouch", "Table"],
    ], [2700, 4860, 1800])
    add_paragraph(doc, "The current proof rebuilds small marts because the data volume is tiny. At larger volume these tables should become dbt incremental models with a merge strategy and explicit unique keys. The staging views can remain views when freshness matters more than repeated query cost.")

    add_heading(doc, "Testing and final audit")
    add_figure(doc, evidence, "Evidence captured from repository tests, AWS and Snowflake", width=6.2)
    add_table(doc, ["Control", "Result", "Evidence"], [
        ["Generator tests", "Passed", "Manifest rows, hashes, reproducibility and synthetic identity checks"],
        ["Airflow structure tests", "Passed", "DAG syntax, six sources, manual date and loopback binding"],
        ["Terraform format", "Passed", "Recursive format check completed"],
        ["Airflow health", "Passed", "HTTP 200 with healthy metadata database, scheduler and triggerer"],
        ["EC2", "Verified", "Instance running with the Airflow instance profile"],
        ["S3", "Verified", "Five manifests, 55 raw objects and about 8.6 MB"],
        ["Snowflake incremental", "Verified", "Five RAW nurse batches and 500 current staged nurses"],
        ["Repository secret check", "Passed", "No supplied Snowflake password or private key marker found in tracked files"],
    ], [2200, 1600, 5560], status_col=1)

    add_heading(doc, "Security and governance decisions")
    add_bullets(doc, [
        "The S3 bucket blocks public access, uses encryption and keeps object versions.",
        "EC2 reaches S3 through an IAM instance profile instead of stored access keys.",
        "Airflow port 8080 binds to loopback and is reached through Systems Manager port forwarding.",
        "Snowflake reads only the approved S3 prefix through a storage integration.",
        "Hightouch and Fivetran use separate RSA service users and least privilege roles.",
        "Customer facing tools receive a governed analytics model rather than unrestricted RAW tables.",
        "The dataset is synthetic and carries no real patient or employee information.",
    ])
    add_callout(doc, "Credential action", "The Snowflake password shared during setup should be rotated. Production access should use key pair authentication, workload identity or a managed secret store. Plaintext secrets must not be committed or pasted into runbooks.", fill=PALE_AMBER, color=AMBER)

    add_heading(doc, "Connector completion status")
    add_paragraph(doc, "The architecture contains two different connector directions. Hightouch sends governed Snowflake data outward. Fivetran brings SaaS data into Snowflake. They should not be presented as interchangeable tools.")
    add_heading(doc, "Hightouch activation path", 2)
    add_table(doc, ["Item", "Status", "Required finish"], [
        ["Snowflake source", "Verified complete", "No action. Four connection tests passed."],
        ["At risk nurse model", "Prepared", "Create or confirm model with NURSE ID as primary key."],
        ["Slack", "Authorization required", "Workspace owner approves OAuth, selects a demo channel and runs initial sync."],
        ["Salesforce", "Planned", "Authorize a sandbox, choose object and map stable identifiers."],
        ["OneDrive", "Planned", "Authorize Microsoft account and choose a controlled export folder."],
        ["Marketo", "Planned", "Authorize destination account and choose person or list operation."],
        ["Pendo", "Planned", "Confirm supported destination action and subscription capability."],
        ["Google Ads", "Planned", "Authorize an ads account and use a consent approved audience."],
        ["Meta Ads", "Planned", "Authorize Business Manager and use a consent approved audience."],
    ], [2200, 2200, 4960], status_col=1)
    add_heading(doc, "Fivetran ingestion path", 2)
    add_table(doc, ["Item", "Status", "Required finish"], [
        ["Snowflake destination", "Verified complete", "Six destination tests passed."],
        ["Marketo source", "Authorization required", "REST endpoint, identity endpoint, client ID and client secret."],
        ["Pendo source", "Authorization required", "Integration key or enabled Data Sync capability."],
        ["SurveyMonkey source", "Authorization required", "OAuth and an active plan that permits connector access."],
        ["Salesforce source", "Authorization required", "Salesforce OAuth and object selection."],
    ], [2200, 2400, 4760], status_col=1)
    add_callout(doc, "Truthful completion rule", "A connector is complete only after authentication, source or destination tests, an initial run, a second incremental run and row level evidence in the receiving system. A saved draft is not a completed integration.", fill=PALE_AMBER, color=AMBER)

    doc.add_page_break()
    add_heading(doc, "System design decisions and tradeoffs")
    add_heading(doc, "Why Snowflake instead of Databricks", 2)
    add_table(doc, ["Decision point", "Snowflake fit", "Databricks fit", "CareMatch choice"], [
        ["Main workload", "SQL analytics, governed sharing and SaaS connectivity", "Large scale Spark, streaming, data engineering and machine learning", "Snowflake"],
        ["Team skill", "Strong SQL and dbt workflow", "Benefits from Spark, Python and platform engineering skills", "Snowflake"],
        ["Data size", "Small structured and semi structured proof", "Best value appears with heavier lakehouse workloads", "Snowflake"],
        ["Operations", "Independent managed compute warehouses", "More control over lakehouse compute and Delta patterns", "Snowflake"],
        ["Future machine learning", "Snowpark is available but not required here", "Strong choice for feature engineering and distributed training", "Reconsider if ML becomes central"],
    ], [1800, 2520, 2720, 2320])
    add_paragraph(doc, "Databricks is not a poor option. Its lakehouse and medallion pattern could represent the same RAW, STAGING and ANALYTICS progression. Snowflake was selected because this proof is SQL first, small and centered on warehouse based activation. Adding Spark would create more platform surface than the present workload needs.")

    add_heading(doc, "Why EC2 Airflow instead of Amazon MWAA", 2)
    add_table(doc, ["EC2 Airflow benefit", "EC2 Airflow cost", "When to move to MWAA"], [
        ["Low proof cost and full control", "Single host failure domain", "Multiple teams depend on uptime"],
        ["Easy to show Docker services", "Team owns upgrades and patching", "Managed patching becomes valuable"],
        ["Simple local SSM tunnel", "No automatic worker scaling", "Queue depth or parallel work grows"],
        ["Exact package control", "Metadata database shares the host", "High availability and managed metadata are required"],
    ], [3120, 3120, 3120])

    add_heading(doc, "Why S3 between Airflow and Snowflake", 2)
    add_paragraph(doc, "S3 separates source execution from warehouse availability. It creates a replay point, preserves source files and allows Snowflake credentials to stay out of Airflow. The tradeoff is an extra storage layer and one more load step. Direct inserts could be simpler for tiny data but would reduce replayability and increase coupling.")
    add_heading(doc, "Why dbt", 2)
    add_paragraph(doc, "dbt keeps transformation SQL beside tests and documentation. It lets Snowflake do the compute and gives reviewers a clear model graph. The tradeoff is that dbt is not an ingestion engine and it does not replace Airflow scheduling, Fivetran connectors or Hightouch delivery.")
    add_heading(doc, "Why Fivetran and Hightouch instead of custom code", 2)
    add_paragraph(doc, "Fivetran owns changing source APIs and ingestion cursors. Hightouch owns destination mappings and warehouse result change detection. This saves integration engineering time. The tradeoffs are subscription cost, vendor limits, account permissions and less control over every API request. High volume or unusual contracts may justify custom services.")

    add_heading(doc, "Likely review questions")
    qa = [
        ("Why does RAW grow while staging stays at 500 nurses?", "RAW stores every dated snapshot. Staging keeps the newest record for each nurse ID."),
        ("Is the pipeline truly incremental?", "Yes at the file and partition level. New filenames load once. A repeated COPY adds zero rows. The source generator creates full daily snapshots."),
        ("Why are there 14 generated files but 11 Snowflake tables?", "Some generated files share one logical table and the file count does not equal the entity table count."),
        ("Why are dbt models visible as views or tables in Snowflake?", "dbt compiles SQL then creates database relations. Staging uses views and marts use tables."),
        ("Why not expose Airflow publicly?", "The case study uses Systems Manager port forwarding so port 8080 does not need a public inbound rule."),
        ("What happens if a task fails?", "Airflow retries. The manifest is not written until all six source upload branches succeed."),
        ("How is a duplicate load prevented?", "Snowflake records loaded filenames. The COPY statements do not use FORCE."),
        ("Can the system recover from bad transformations?", "RAW files remain in S3 and RAW history stays available. Staging and marts can be rebuilt."),
        ("Why not use one service for everything?", "Orchestration, durable storage, analytics, transformation and activation have different operational needs."),
        ("Is Slack already receiving data?", "No. The destination is selected but workspace OAuth and a verified sync are still required."),
        ("Is Fivetran already ingesting Marketo?", "No. The Snowflake destination is tested but Marketo API credentials are still required."),
        ("What is the largest production risk?", "The single EC2 Airflow host and manual Snowflake load step are the largest reliability gaps."),
        ("What would change first for production?", "Move orchestration to a managed highly available service, automate Snowflake loading, centralize secrets and add monitoring alerts."),
    ]
    for question, answer in qa:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        set_run(p.add_run(question + " "), bold=True, color=NAVY)
        set_run(p.add_run(answer))

    add_heading(doc, "Production readiness roadmap")
    add_table(doc, ["Priority", "Work", "Acceptance evidence"], [
        ["1", "Rotate exposed credentials and centralize secrets", "No plaintext secrets and successful key based access"],
        ["2", "Automate S3 to Snowflake load with Snowpipe or an Airflow task", "A new S3 partition reaches RAW without a manual worksheet"],
        ["3", "Move Airflow metadata and workers to a highly available managed design", "Failure of one compute host does not stop scheduling"],
        ["4", "Complete Slack OAuth and run two Hightouch syncs", "Initial operations plus a later change only sync"],
        ["5", "Complete one Fivetran SaaS source", "Initial history plus a later cursor based incremental sync"],
        ["6", "Add monitoring and cost alerts", "Alerts for failed DAGs, stale data, warehouse spend and connector failures"],
        ["7", "Add environment separation and CI", "Plan, test and dbt checks run before deployment"],
    ], [950, 4250, 4160])

    demo_heading = add_heading(doc, "Demonstration script")
    demo_heading.paragraph_format.page_break_before = True
    add_numbered(doc, [
        "Start with the architecture visual and explain the two branches.",
        "Open Airflow and show the eight successful tasks for the latest date.",
        "Open S3 and show the matching date partition plus manifest.",
        "Run the RAW batch date query in Snowflake and show five dates.",
        "Compare RAW nurses with staging nurses and explain deduplication.",
        "Open the dbt model files and match them to Snowflake relations.",
        "Run the data quality query and require zero failing rows.",
        "Show the Hightouch source tests and state that Slack OAuth is pending.",
        "Show the Fivetran destination tests and state that SaaS source credentials are pending.",
        "End with the production roadmap and answer the Snowflake versus Databricks question.",
    ])

    add_heading(doc, "Repository map")
    add_table(doc, ["Area", "Location"], [
        ["Synthetic data generator", "src/generate_healthcare_data.py"],
        ["Airflow DAG", "airflow/dags/synthetic_sources_to_s3.py"],
        ["Terraform", "infra/terraform"],
        ["Snowflake SQL", "snowflake/sql"],
        ["dbt project", "dbt/models and dbt/tests"],
        ["Runbooks and evidence", "docs"],
        ["Automated tests", "tests"],
    ], [3000, 6360])

    add_heading(doc, "Sources and evidence")
    add_bullets(doc, [
        "Snowflake customer case study: IntelyCare grows during COVID and saves marketing spend with Hightouch, Snowflake and Fivetran. Supplied PDF, 2022.",
        "CareMatch Git repository and runbooks. Local main branch audited on 27 August 2026.",
        "Read only AWS audit: EC2 instance state, S3 manifests and S3 raw object listing, 27 August 2026.",
        "Snowflake worksheet evidence: five RAW nurse dates, 2,500 RAW nurse rows and 500 staged nurses, 27 August 2026.",
        "Snowflake documentation: key concepts, virtual warehouses and COPY based ingestion.",
        "Databricks documentation: lakehouse and medallion architecture.",
        "AWS documentation: Amazon Managed Workflows for Apache Airflow architecture and operations.",
        "Fivetran documentation: initial sync, incremental sync and connector cursors.",
        "Hightouch documentation: sync modes and difference based change data capture.",
    ])
    add_paragraph(doc, "Official references: https://docs.snowflake.com/en/user-guide/intro-key-concepts | https://docs.databricks.com/aws/en/lakehouse/medallion | https://docs.aws.amazon.com/mwaa/ | https://fivetran.com/docs/getting-started/glossary | https://hightouch.com/docs/syncs/cdc", italic=True)

    doc.core_properties.title = "CareMatch Modern Data Stack Case Study"
    doc.core_properties.subject = "Healthcare staffing data platform architecture and implementation"
    doc.core_properties.author = "CareMatch Data Platform"
    doc.core_properties.keywords = "Airflow, AWS, S3, Snowflake, dbt, Hightouch, Fivetran"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
